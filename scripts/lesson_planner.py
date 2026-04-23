from __future__ import annotations

from pathlib import Path

import nbformat
from docx import Document
from docx.shared import Pt


def validate_lesson_plan(content: dict, good_lesson: dict) -> None:
    required_top_level = [
        "lesson_flow",
        "classroom_timing",
        "misconceptions",
        "in_class_checks",
        "teacher_quality_review",
    ]
    required_lesson_fields = [
        "teacher_verdict",
        "fixed_template",
    ]
    missing = [key for key in required_top_level if key not in content]
    missing.extend([f"lesson.{key}" for key in required_lesson_fields if key not in content.get("lesson", {})])
    if missing:
        raise ValueError(f"Lesson plan is missing fixed-template fields: {', '.join(missing)}")

    expected_template = good_lesson["fixed_lesson_template"]
    actual_template = content["lesson"]["fixed_template"]
    if len(actual_template) != len(expected_template):
        raise ValueError("Lesson fixed template must match good_lesson.fixed_lesson_template length.")
    for actual, expected in zip(actual_template, expected_template):
        if not expected.startswith(actual):
            raise ValueError(f"Lesson fixed template item does not match quality config: {actual}")


def configure_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)


def add_doc_bullets(doc: Document, bullets: list[str]) -> None:
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def build_docx(course: dict, chapter: dict, content: dict, good_lesson: dict, teacher_profile: str, output_path: Path) -> None:
    doc = Document()
    configure_doc_styles(doc)

    lesson = content["lesson"]
    doc.add_heading(f"{lesson['title']}：教学包", level=0)
    doc.add_paragraph(f"课程：{course['course']['name']}")
    doc.add_paragraph(f"对象：{lesson['audience']}")
    doc.add_paragraph(f"课时：{lesson['duration_minutes']} 分钟")
    doc.add_paragraph(f"核心信息：{lesson['key_message']}")

    verdict = lesson["teacher_verdict"]
    doc.add_heading("讲前判断：我会不会拿去上课", level=1)
    doc.add_paragraph(f"结论：{verdict['verdict']}")
    doc.add_paragraph(f"理由：{verdict['reason']}")

    doc.add_heading("什么叫好课", level=1)
    doc.add_paragraph(good_lesson["good_lesson"]["teacher_verdict_standard"])
    for item in good_lesson["good_lesson"]["definition"]:
        doc.add_paragraph(f"{item['name']}：{item['requirement']} 证据：{item['evidence']}", style="List Bullet")

    doc.add_heading("固定教学模板", level=1)
    add_doc_bullets(doc, good_lesson["fixed_lesson_template"])

    doc.add_heading("课堂内容顺序", level=1)
    for item in content["lesson_flow"]:
        doc.add_paragraph(f"{item['stage']}：{item['purpose']}", style="List Number")
        doc.add_paragraph(f"核心内容：{item['core_content']}")

    doc.add_heading("学习目标", level=1)
    add_doc_bullets(doc, lesson["learning_outcomes"])

    doc.add_heading("时间分配与理解证据", level=1)
    timing_table = doc.add_table(rows=1, cols=4)
    timing_table.style = "Table Grid"
    for index, header in enumerate(["环节", "分钟", "教师动作", "学生理解证据"]):
        timing_table.rows[0].cells[index].text = header
    for item in content["classroom_timing"]:
        row = timing_table.add_row().cells
        row[0].text = item["stage"]
        row[1].text = str(item["minutes"])
        row[2].text = item["teacher_action"]
        row[3].text = item["student_evidence"]

    sections = content["doc_sections"]
    for key in ("teaching_plan", "lecture_script"):
        section = sections[key]
        doc.add_heading(section["title"], level=1)
        for paragraph in section["paragraphs"]:
            doc.add_paragraph(paragraph)

    doc.add_heading(sections["blackboard_plan"]["title"], level=1)
    add_doc_bullets(doc, sections["blackboard_plan"]["bullets"])

    doc.add_heading(sections["classroom_questions"]["title"], level=1)
    for item in sections["classroom_questions"]["items"]:
        doc.add_paragraph(f"问：{item['question']}", style="List Number")
        doc.add_paragraph(f"答：{item['answer']}")

    doc.add_heading("常见误区与教师处理", level=1)
    misconception_table = doc.add_table(rows=1, cols=3)
    misconception_table.style = "Table Grid"
    for index, header in enumerate(["误区", "纠正", "教师动作"]):
        misconception_table.rows[0].cells[index].text = header
    for item in content["misconceptions"]:
        row = misconception_table.add_row().cells
        row[0].text = item["misconception"]
        row[1].text = item["correction"]
        row[2].text = item["teacher_move"]

    doc.add_heading("课堂检查", level=1)
    for item in content["in_class_checks"]:
        doc.add_paragraph(f"{item['check']}：{item['prompt']}", style="List Number")
        doc.add_paragraph(f"期待回答：{item['expected']}")

    doc.add_heading("作业、答案与评分 Rubrics", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["类型", "题目", "参考答案", "评分标准"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for item in content["homework"]:
        row = table.add_row().cells
        row[0].text = item["type"]
        row[1].text = item["question"]
        row[2].text = item["answer"]
        row[3].text = item["rubric"]

    doc.add_heading("教师画像摘要", level=1)
    doc.add_paragraph(teacher_profile.strip()[:900])

    doc.add_heading("章节配置摘要", level=1)
    doc.add_paragraph(f"AI 应用入口：{chapter['chapter']['ai_application_entry']}")
    doc.add_paragraph(f"实验目标：{chapter['chapter']['lab_goal']}")
    doc.add_paragraph(f"教师拿走的一句话：{chapter['chapter']['teacher_takeaway']}")

    quality = content["teacher_quality_review"]
    doc.add_heading("课后质量自查", level=1)
    doc.add_paragraph(f"总分：{quality['total_score']} / 100，阈值：{quality['pass_threshold']}")
    doc.add_paragraph(f"结论：{quality['verdict']}")
    doc.add_paragraph("风险：")
    add_doc_bullets(doc, quality["risks"])
    doc.add_paragraph("修复动作：")
    add_doc_bullets(doc, quality["repair_actions"])

    doc.save(output_path)


def build_notebook(content: dict, output_path: Path) -> None:
    lesson = content["lesson"]
    nb = nbformat.v4.new_notebook()
    nb["cells"] = [
        nbformat.v4.new_markdown_cell(
            f"# {lesson['title']}实验\n\n"
            "目标：通过一维二次损失函数观察梯度下降、学习率和损失下降之间的关系。"
        ),
        nbformat.v4.new_markdown_cell(
            "## 实验问题\n\n"
            "- 学习率过小会发生什么？\n"
            "- 学习率合适时，参数如何接近最优点？\n"
            "- 学习率过大时，为什么可能震荡或发散？"
        ),
        nbformat.v4.new_code_cell(
            "import numpy as np\n"
            "import matplotlib.pyplot as plt\n"
            "plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'Arial Unicode MS', 'DejaVu Sans']\n"
            "plt.rcParams['axes.unicode_minus'] = False\n"
            "\n"
            "def loss(w):\n"
            "    return (w - 3) ** 2 + 1\n"
            "\n"
            "def grad(w):\n"
            "    return 2 * (w - 3)\n"
            "\n"
            "def run_descent(eta, steps=25, w0=0.0):\n"
            "    ws = [w0]\n"
            "    losses = [loss(w0)]\n"
            "    w = w0\n"
            "    for _ in range(steps):\n"
            "        w = w - eta * grad(w)\n"
            "        ws.append(w)\n"
            "        losses.append(loss(w))\n"
            "    return np.array(ws), np.array(losses)\n"
        ),
        nbformat.v4.new_code_cell(
            "learning_rates = [0.05, 0.2, 1.05]\n"
            "results = {eta: run_descent(eta) for eta in learning_rates}\n"
            "\n"
            "fig, axes = plt.subplots(1, 2, figsize=(12, 4))\n"
            "x = np.linspace(-2, 8, 300)\n"
            "axes[0].plot(x, loss(x), color='black', linewidth=2, label='L(w)')\n"
            "for eta, (ws, losses) in results.items():\n"
            "    axes[0].plot(ws, loss(ws), marker='o', label=f'eta={eta}')\n"
            "    axes[1].plot(losses, marker='o', label=f'eta={eta}')\n"
            "\n"
            "axes[0].set_title('参数在损失曲线上的轨迹')\n"
            "axes[0].set_xlabel('w')\n"
            "axes[0].set_ylabel('L(w)')\n"
            "axes[0].legend()\n"
            "axes[1].set_title('每一步的损失变化')\n"
            "axes[1].set_xlabel('step')\n"
            "axes[1].set_ylabel('loss')\n"
            "axes[1].legend()\n"
            "plt.tight_layout()\n"
            "plt.savefig('loss_trajectories.png', dpi=160)\n"
        ),
        nbformat.v4.new_markdown_cell(
            "## 观察记录\n\n"
            "1. `eta = 0.05`：通常下降稳定但较慢。\n"
            "2. `eta = 0.2`：较快接近最低点。\n"
            "3. `eta = 1.05`：可能出现震荡或不稳定。\n\n"
            "把这些现象和更新公式 `w_{t+1} = w_t - eta * dL/dw` 对应起来。"
        ),
    ]
    nbformat.write(nb, output_path)


def write_quality_report(course: dict, chapter: dict, content: dict, good_lesson: dict, output_path: Path) -> None:
    review = content["quality_review"]
    lesson = content["lesson"]
    teacher_review = content["teacher_quality_review"]
    lines = [
        f"# {lesson['title']}教学质量门禁报告",
        "",
        "## 我会不会拿去上课",
        "",
        f"- 结论：{teacher_review['verdict']}",
        f"- 总分：{teacher_review['total_score']} / 100",
        f"- 通过阈值：{teacher_review['pass_threshold']}",
        "",
        "## 什么叫好课",
        "",
        f"- 判断标准：{good_lesson['good_lesson']['teacher_verdict_standard']}",
        f"- 最低结论：{good_lesson['good_lesson']['minimum_verdict']}",
        "",
    ]
    for item in good_lesson["good_lesson"]["definition"]:
        lines.append(f"- {item['name']}：{item['requirement']}")
    lines.extend([
        "",
        "## 三重质量门禁",
        "",
        f"- 教学可讲：{review['teaching']}",
        f"- 数学可信：{review['math']}",
        f"- AI 连接有效：{review['ai_connection']}",
        "",
        "## 每章必须回答的四个问题",
        "",
        "1. 学生为什么需要学这个数学概念？",
        "   - 因为梯度说明误差如何转化为参数更新方向。",
        "2. 这个概念解释了 AI 模型里的哪个关键设计？",
        "   - 解释训练循环中的损失计算、梯度求解和参数更新。",
        "3. 学生如何通过实验看到这个数学概念在起作用？",
        "   - 通过 notebook 对比不同学习率下的损失曲线和参数轨迹。",
        "4. 老师如何判断学生是否真正理解？",
        "   - 看学生能否解释梯度方向、学习率影响和实验曲线含义。",
        "",
        "## 逐项评分",
        "",
    ])
    for name, score in teacher_review["scores"].items():
        max_score = good_lesson["quality_score"]["dimensions"][name]
        lines.append(f"- {name}: {score} / {max_score}")
    lines.extend([
        "",
        "## 风险与修复动作",
        "",
        "### 风险",
        "",
    ])
    lines.extend([f"- {item}" for item in teacher_review["risks"]])
    lines.extend([
        "",
        "### 修复动作",
        "",
    ])
    lines.extend([f"- {item}" for item in teacher_review["repair_actions"]])
    lines.extend([
        "",
        "## 生成依据",
        "",
        f"- 课程：{course['course']['name']}",
        f"- 章节：{chapter['chapter']['title']}",
        "- 好课标准：configs/quality/good_lesson.yaml",
        "- Lesson plan：sources/chapters/gradient_descent/content.yaml",
    ])
    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
